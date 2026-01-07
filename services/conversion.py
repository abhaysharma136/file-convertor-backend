from reportlab.pdfgen import canvas
from pypdf import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter

def docx_to_pdf(input_path: str, output_path: str):
    doc = Document(input_path)
    pdf = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    y_position = height - 40
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pdf.drawString(40, y_position, text)
            y_position -= 14

            if y_position < 40:
                pdf.showPage()
                y_position = height - 40

    pdf.save()

def pdf_to_docx(input_path: str, output_path: str):
    reader = PdfReader(input_path)
    doc = Document()

    for page in reader.pages:
        text = page.extract_text()
        if text:
            for line in text.split("\n"):
                doc.add_paragraph(line)

    doc.save(output_path)

def convert_document(input_path: str, output_path: str, target_format: str):
    if target_format == "pdf":
        docx_to_pdf(input_path, output_path)
    elif target_format == "docx":
        pdf_to_docx(input_path, output_path)
    else:
        raise ValueError("Unsupported document format")